# -*- coding: utf-8 -*-
"""build_report.py — assembla la RP 8.6 (DOCX) sul template ufficiale START.

Apre il template `RPX.Y Titolo_Relazione_Parziale_data.docx` (intestazioni, loghi,
piè di pagina e stili), compila i segnaposto di copertina, elimina il corpo di
esempio da "1.INTRODUZIONE" in poi e ricostruisce il contenuto della relazione
(testo, tabelle, figure). I valori numerici sono importati dal motore `impatto.py`
per garantire la coerenza tra codice, CSV, figure e documento.

Stile coerente con la RP 7.3 (verde START, header di tabella verde scuro).

Uso:  dalla cartella RP8.6/  ->  python -m src.build_report
      (oppure)               ->  python src/build_report.py
"""
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import impatto  # noqa: E402

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))     # .../RP8.6
TPL = os.path.join(BASE, "..", "RPX.Y Titolo_Relazione_Parziale_data.docx")
FIG = os.path.join(BASE, "docs", "figures")
OUTDOCX = os.path.join(BASE, "docs", "RP8.6_Assessment_Impatto.docx")
TODAY = "07.08.2026"

DGREEN = RGBColor(0x02, 0x4C, 0x41)
GREY = RGBColor(0x33, 0x33, 0x33)
HDR_FILL = "0B5A3C"


# ---- formattazione numeri (stile italiano: virgola decimale) ---------------------
def n1(x):
    return f"{x:,.1f}".replace(",", "§").replace(".", ",").replace("§", ".")


def n2(x):
    return f"{x:.2f}".replace(".", ",")


def n0(x):
    return f"{x:,.0f}".replace(",", ".")


# ---- helper di layout ------------------------------------------------------------
def set_para_text(p, text):
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


def h1(doc, t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DGREEN
    r.font.name = "Calibri"
    return p


def h2(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = DGREEN
    r.font.name = "Calibri"
    return p


def _inline(p, text, size=11):
    for i, seg in enumerate(text.split("**")):
        if seg == "":
            continue
        r = p.add_run(seg)
        r.font.size = Pt(size)
        r.font.name = "Calibri"
        r.bold = (i % 2 == 1)


def para(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    _inline(p, text, size)
    return p


def meta(doc, key, val):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(key + " ")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Calibri"
    r.font.color.rgb = DGREEN
    r2 = p.add_run(val)
    r2.font.size = Pt(10.5)
    r2.font.name = "Calibri"
    return p


def bullet(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(2)
    p.add_run("•  ").font.size = Pt(11)
    _inline(p, text, 11)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    return p


def note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    _inline(p, text, 10)
    for r in p.runs:
        r.italic = True
        r.font.color.rgb = GREY
    return p


def figure(doc, name, width_mm=150):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    png = os.path.join(FIG, name + ".png")
    p.add_run().add_picture(png, width=Mm(width_mm))
    return p


def _shade(tc, hexc):
    tcPr = tc._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def _border(tc):
    tcPr = tc._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "BFBFBF")
        b.append(e)
    tcPr.append(b)


def table(doc, headers, rows, right_from=1, fs=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hd = t.rows[0]
    trPr = hd._tr.get_or_add_trPr()
    e = OxmlElement("w:tblHeader")
    e.set(qn("w:val"), "true")
    trPr.append(e)
    for i, htext in enumerate(headers):
        c = hd.cells[i]
        c.text = ""
        _shade(c, HDR_FILL)
        _border(c)
        pr = c.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pr.add_run(htext)
        r.bold = True
        r.font.size = Pt(fs)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Calibri"
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            c.text = ""
            _border(c)
            pr = c.paragraphs[0]
            s = str(val)
            bold = s.startswith("**")
            s = s.replace("**", "")
            pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= right_from else WD_ALIGN_PARAGRAPH.LEFT
            r = pr.add_run(s)
            r.font.size = Pt(fs)
            r.font.name = "Calibri"
            r.bold = bold
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# =======================================================================
def build():
    m = impatto.multipliers()
    imps = {r: impatto.impact(r) for r in impatto.FINAL_DEMAND}
    s = impatto.summary()

    doc = Document(TPL)
    cover = {
        "Titolo relazione": "ASSESSMENT DELL'IMPATTO DI START",
        "Relazione Parziale N°: RPX.Y": "Relazione Parziale N°: RP 8.6",
        "Versione del Documento: RV.X": "Versione del Documento: V1.0",
        "Data di Revisione del Documento: XX.YY.ZZ": f"Data di Revisione del Documento: {TODAY}",
        "Responsabilità:  Partner - Ruolo": "Responsabilità: Gresmalt - Capofila (in collaborazione con le Università)",
    }
    intro_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        for k, v in cover.items():
            if t == k.strip():
                set_para_text(p, v)
        if t.upper().replace(" ", "").startswith("1.INTRODUZIONE"):
            intro_idx = i
    if intro_idx is not None:
        for p in doc.paragraphs[intro_idx:]:
            p._element.getparent().remove(p._element)

    # --- blocco identificativo (scheda) ---
    meta(doc, "Progetto:", "START — SusTainable dAta-dRiven manufacTuring (DM 31/12/2021, Accordi per l'Innovazione — MIMIT)")
    meta(doc, "Obiettivo Realizzativo:", "OR 8 — Project management, misurazione dei risultati e analisi degli scostamenti (Sviluppo Sperimentale)")
    meta(doc, "Attività:", "8.6 — Assessment dell'impatto di START")
    meta(doc, "Risultato parziale / KPI:", "Report sulla valutazione d'impatto — Analisi / Matrice input-output (baseline: Nessuna informazione → obiettivo: Analisi effettuata)")
    meta(doc, "Risultato finale di OR:", "RF 8 — Piano di Coordinamento del Progetto START")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    note(doc, "Natura di questa relazione. In assenza dei dati macroeconomici regionali (problema "
              "progettuale n. 6), l'attività 8.6 è impostata come simulazione dell'impatto potenziale e, "
              "soprattutto, come predisposizione e validazione del modello analitico input-output da "
              "applicare ex-post — nei mesi e negli anni successivi alla chiusura del progetto — con i "
              "consuntivi degli investimenti e le tavole input-output regionali ISTAT. Il deliverable è "
              "quindi lo strumento (matrice/modello riproducibile + protocollo di misurazione); i valori "
              "riportati sono l'esito di una simulazione con dati rappresentativi e servono a dimostrare "
              "l'operatività del modello e l'ordine di grandezza degli effetti attesi, non costituiscono "
              "una misura consuntiva dell'impatto.")

    # ============= 1. INTRODUZIONE =============
    h1(doc, "1. INTRODUZIONE")
    h2(doc, "1.1 Background")
    para(doc, "L'OR 8 attribuisce a START una finalità dichiarata come **inedita per i progetti di R&S a "
              "carattere industriale**: la **valutazione dell'impatto del progetto sulla società di "
              "riferimento e sui territori**. L'attività 8.6, svolta in collaborazione con le Università "
              "del partenariato, ne predispone lo strumento: un **modello di analisi input-output** — sul "
              "modello econometrico di **Wassily Leontief** — capace di stimare l'impatto diretto, "
              "indiretto e indotto dei risultati del progetto sui quattro territori dove operano le "
              "attività industriali e di ricerca: l'**Emilia-Romagna** (distretto ceramico di Sassuolo, "
              "sede di Gresmalt e SACMI), la **Provincia di Bolzano** (Libera Università di Bolzano — "
              "Digital Twin e AI), la **Sardegna** (Università di Sassari — Architectural Design 4.0+ e "
              "involucro edilizio ceramico) e la **Calabria** (Università della Calabria — modelli ML/ANN "
              "e diagnostica non distruttiva). Poiché la misura consuntiva richiede dati non ancora "
              "disponibili, l'attività si concretizza in una **simulazione** che valorizza il modello e "
              "ne dimostra l'applicazione, lasciandolo pronto per la **misurazione ex-post**.")
    para(doc, "I risultati di START — la transizione *Smart Factory → Intelligent Factory → Intelligent "
              "Industry* guidata dall'AI, le quattro impronte di sostenibilità e la contabilità "
              "exergetica EEA+, l'architettura Edge-to-Cloud, il data-driven product design, il collaudo "
              "in ambiente operativo — attivano produzione, valore aggiunto e occupazione non solo "
              "nell'impresa focale ma lungo tutta la filiera locale (estrazione, chimica, logistica, "
              "energia, servizi/R&S) e attraverso i consumi indotti dai redditi generati. L'analisi "
              "input-output è lo strumento elettivo per **catturare questi effetti a cascata**.")
    h2(doc, "1.2 Il problema dei dati, la scelta della simulazione e il legame con la RP 8.5")
    para(doc, "L'analisi input-output di impatto richiede **dati macroeconomici su scala regionale** "
              "(tavole input-output/SUT, coefficienti settoriali) e **consuntivi degli investimenti** del "
              "progetto, **non direttamente disponibili** al Team durante la vita del progetto (**problema "
              "progettuale n. 6**). Per questo l'attività 8.6 non pretende una misura definitiva "
              "dell'impatto, ma: (1) **imposta e valida il modello analitico** parametrizzato sulla "
              "filiera ceramica; (2) lo **esegue in simulazione** con coefficienti rappresentativi; (3) "
              "definisce il **protocollo di applicazione ex-post** (§ 2.5). La soluzione prevista in "
              "scheda per l'accesso ai dati è, in accordo con il **Piano di Stakeholder Engagement (RP "
              "8.5)**, il **coinvolgimento delle autorità territoriali** (ISTAT, uffici statistici "
              "regionali e provinciali): **RP 8.5 è quindi il presupposto della misurazione ex-post** qui "
              "predisposta.")
    h2(doc, "1.3 Baseline")
    para(doc, "La **baseline** del KPI è **«Nessuna informazione»**: prima dell'attività 8.6 non esiste "
              "un modello di analisi input-output dell'impatto territoriale del progetto. L'**obiettivo** "
              "è l'**analisi effettuata**, intesa qui come il **modello (matrice I-O) impostato ed "
              "eseguito in simulazione** sulle quattro regioni e pronto per l'applicazione ex-post: lo "
              "strumento è il risultato, il dato consuntivo verrà dopo.")
    h2(doc, "1.4 Scopi dell'attività")
    bullet(doc, "**impostare e validare** il modello input-output di Leontief per la filiera ceramica "
                "(coefficienti tecnici, inversa di Leontief, moltiplicatori Tipo I/II);")
    bullet(doc, "**simulare** l'impatto diretto, indiretto e indotto del progetto sulle quattro regioni "
                "in termini di output, valore aggiunto (contributo al PIL) e occupazione;")
    bullet(doc, "**definire il protocollo di applicazione ex-post** per ri-eseguire il modello con le "
                "tavole I-O regionali ISTAT e i consuntivi (via RP 8.5), alimentando l'analisi degli "
                "scostamenti dell'OR 8;")
    bullet(doc, "**produrre il Report di valutazione d'impatto** (KPI: analisi effettuata) come strumento "
                "riproducibile del Piano di Coordinamento (RF 8).")

    # ============= 2. METODOLOGIA =============
    h1(doc, "2. METODOLOGIA")
    h2(doc, "2.1 Il modello input-output di Leontief")
    para(doc, "L'analisi input-output rappresenta le interdipendenze tra i settori di un'economia. Data "
              "la **matrice dei coefficienti tecnici** A, dove a(i,j) è l'input dal settore i necessario "
              "per produrre un'unità di output del settore j, la produzione totale x che soddisfa una "
              "domanda finale f è x = A·x + f, da cui **x = (I − A)⁻¹ · f = L · f**, dove **L = (I − A)⁻¹** "
              "è la **matrice inversa di Leontief**. Applicata a una variazione della domanda finale Δf "
              "(lo shock del progetto), fornisce la variazione della produzione totale **Δx = L·Δf**, che "
              "incorpora gli effetti a cascata lungo la filiera (Leontief, 1936; Miller & Blair, 2009). "
              "Lo schema del modello è in Figura 1.")
    figure(doc, "fig_imp1_metodo", 165)
    caption(doc, "Figura 1 — Dalla domanda finale del progetto, tramite l'inversa di Leontief, all'output "
                 "attivato e alle sue dimensioni di impatto (valore aggiunto, occupazione).")
    h2(doc, "2.2 Effetti diretti, indiretti e indotti; moltiplicatori Tipo I e Tipo II")
    para(doc, "L'impatto si scompone in tre componenti: **diretto** (la domanda finale attivata dal "
              "progetto, Δf); **indiretto** (la produzione attivata lungo la filiera dei fornitori); "
              "**indotto** (la produzione attivata dai consumi delle famiglie finanziati dai redditi da "
              "lavoro generati). Si calcolano due famiglie di **moltiplicatori** di produzione: **Tipo I** "
              "— modello aperto (diretto + indiretto), somma per colonna della matrice L; **Tipo II** — "
              "modello chiuso rispetto alle famiglie (aggiunge l'indotto), con A estesa da una riga "
              "(redditi da lavoro per unità di output) e una colonna (consumi per unità di reddito). "
              "L'impatto su **valore aggiunto** e **occupazione** si ottiene applicando all'output "
              "attivato i rispettivi coefficienti settoriali (quota di VA sull'output; ULA per milione di "
              "euro di output).")
    h2(doc, "2.3 Settori, coefficienti e shock del progetto (parametri di simulazione)")
    para(doc, "Il modello adotta **sei settori** rappresentativi della filiera ceramica e del suo indotto "
              "territoriale: estrazione di minerali non metalliferi; fabbricazione di prodotti ceramici; "
              "chimica (smalti, additivi, chemicals); trasporti e logistica; energia; servizi (ICT, R&S, "
              "professionali — comparto che in START incorpora AI, digital twin e progettazione). Lo "
              "**shock di domanda finale** è **differenziato per regione** secondo il ruolo del soggetto "
              "attuatore: concentrato su ceramica/estrazione/chimica/logistica/energia in "
              "**Emilia-Romagna**; orientato a ricerca, ICT/R&S e servizi in **Provincia di Bolzano** e "
              "**Calabria**; orientato a servizi/progettazione e applicazione del prodotto ceramico in "
              "**Sardegna**. La **struttura** del modello coincide con quella del progetto gemello VOLT "
              "(RP 9.6) — stessa industria e stesso distretto; ciò che è stato ricalibrato per START è la "
              "**ripartizione regionale** dello shock sui quattro territori della scheda.")
    h2(doc, "2.4 Trattamento dei dati e riproducibilità")
    para(doc, "Coerentemente con il problema n. 6, i coefficienti della matrice A, i coefficienti di VA e "
              "occupazione e lo shock di domanda finale sono **rappresentativi**, calibrati sull'ordine di "
              "grandezza della filiera; **non contengono dati riservati** e vanno **consolidati** con le "
              "tavole input-output regionali ISTAT (SUT/IO), il cui accesso passa dal coinvolgimento delle "
              "autorità territoriali previsto nella RP 8.5. Il modello è codificato e riproducibile "
              "(`RP8.6/src/impatto.py`, `RP8.6/run_impatto.py`), con export `RP8.6/output/impatto_*.csv`: "
              "ri-eseguirlo con i dati ufficiali è sufficiente a produrre la stima consuntiva senza "
              "modificare la struttura analitica.")
    h2(doc, "2.5 Protocollo di applicazione ex-post")
    para(doc, "Il valore dell'attività 8.6 è il modello e il suo **protocollo di applicazione ex-post**, "
              "da attivare nei mesi e anni successivi alla chiusura del progetto secondo quattro passi: "
              "(1) **acquisizione dati** via autorità territoriali (RP 8.5): tavole I-O regionali ISTAT e "
              "consuntivi degli investimenti; (2) **ricalibrazione** dei coefficienti e dello shock con i "
              "valori ufficiali, a struttura invariata; (3) **ri-esecuzione** del calcolo per la stima "
              "consuntiva; (4) **analisi degli scostamenti** tra simulato (previsione) e consuntivo, che "
              "alimenta il Piano di Coordinamento (RF 8).")
    table(doc,
          ["Elemento", "Ora (simulazione, RP 8.6)", "Ex-post (misurazione)"],
          [["Matrice dei coefficienti tecnici A", "Rappresentativa (filiera ceramica)", "Tavole I-O regionali ISTAT"],
           ["Shock di domanda finale", "Stimato (investimenti + produzione attesa)", "Consuntivi del progetto"],
           ["Coefficienti VA / occupazione", "Rappresentativi", "Contabilità regionale / ISTAT"],
           ["Esito", "Impatto potenziale (ordine di grandezza)", "Impatto consuntivo verificabile"],
           ["Uso", "Validazione del modello, stima attesa", "Analisi degli scostamenti (OR 8)"]],
          right_from=99, fs=8.5)
    caption(doc, "Tabella — Stato attuale (simulazione) vs applicazione ex-post.")

    # ============= 3. RISULTATI =============
    h1(doc, "3. RISULTATI DELLA SIMULAZIONE")
    para(doc, "I risultati che seguono sono l'esito della **simulazione** del modello con parametri "
              "rappresentativi (§ 2.3): dimostrano l'operatività dello strumento e l'ordine di grandezza "
              "degli effetti attesi, e saranno riprodotti con i dati ufficiali secondo il protocollo "
              "ex-post (§ 2.5). I **moltiplicatori** (§ 3.1) sono la componente più robusta del modello, "
              "perché dipendono dalla struttura dei coefficienti tecnici più che dall'entità dello shock.")
    h2(doc, "3.1 Moltiplicatori di produzione della filiera")
    para(doc, f"Il settore **ceramico** ha il moltiplicatore più elevato (**{n2(m['tipo1'][1])}** di Tipo "
              f"I, **{n2(m['tipo2'][1])}** di Tipo II), coerentemente con la sua posizione di cuore della "
              "filiera e con la ricchezza dei legami a monte (estrazione, chimica, energia, logistica); "
              f"seguono chimica ({n2(m['tipo1'][2])}/{n2(m['tipo2'][2])}) ed estrazione "
              f"({n2(m['tipo1'][0])}/{n2(m['tipo2'][0])}). L'inclusione dell'effetto indotto (Tipo II) "
              "accresce sensibilmente i moltiplicatori dei settori a maggiore intensità di lavoro "
              f"(servizi: da {n2(m['tipo1'][5])} a {n2(m['tipo2'][5])}).")
    figure(doc, "fig_imp2_moltiplicatori", 150)
    caption(doc, "Figura 2 — Moltiplicatori di produzione per settore: il settore ceramico traina "
                 "l'attivazione dell'output regionale.")
    table(doc, ["Settore", "Tipo I (aperto)", "Tipo II (chiuso)"],
          [[impatto.SECTORS[j], n2(m["tipo1"][j]), n2(m["tipo2"][j])] for j in range(impatto.N)],
          right_from=1)
    caption(doc, "Tabella 1 — Moltiplicatori di produzione per settore.")
    h2(doc, "3.2 Impatto simulato sull'output per regione")
    er = imps["Emilia-Romagna"]
    para(doc, f"In **Emilia-Romagna** i **{n1(er['output_diretto'])} M€** di domanda finale diretta "
              f"attiverebbero complessivamente **{n1(er['output_totale'])} M€** di output (indiretto "
              f"{n1(er['output_indiretto'])} M€; indotto {n1(er['output_indotto'])} M€), con un "
              f"**moltiplicatore di {n2(er['moltiplicatore_output'])}**. Nelle tre regioni della ricerca "
              "lo shock è più contenuto e orientato a settori con legami a monte meno intensi rispetto al "
              f"core ceramico. Complessivamente, **{n1(s['output_diretto_totale'])} M€** di domanda "
              f"diretta attiverebbero **~{n1(s['output_attivato_totale'])} M€** di output totale "
              f"(moltiplicatore medio **{n2(s['moltiplicatore_medio'])}**).")
    figure(doc, "fig_imp3_regioni", 150)
    caption(doc, "Figura 3 — Simulazione dell'output attivato per regione, scomposto in diretto, "
                 "indiretto (filiera) e indotto (consumi). Valori rappresentativi.")
    h2(doc, "3.3 Impatto simulato su valore aggiunto e occupazione")
    para(doc, "L'output attivato si tradurrebbe in **valore aggiunto** (contributo al PIL regionale) e "
              f"**occupazione** (Figura 4). In totale la simulazione stima **~{n1(s['valore_aggiunto_totale'])} "
              f"M€** di valore aggiunto e **~{n0(s['occupazione_ula_totale'])} ULA** attivati sui quattro "
              "territori.")
    figure(doc, "fig_imp4_impatto", 165)
    caption(doc, "Figura 4 — Simulazione del contributo al valore aggiunto (sinistra) e dell'occupazione "
                 "attivata (destra), scomposti in diretto/indiretto/indotto. Valori rappresentativi.")
    reg_rows = []
    for r, i in imps.items():
        reg_rows.append([r, n1(i["output_diretto"]), n1(i["output_totale"]),
                         n2(i["moltiplicatore_output"]), n1(i["va_totale"]), n0(i["occ_totale"])])
    reg_rows.append(["**Totale**", "**" + n1(s["output_diretto_totale"]) + "**",
                     "**" + n1(s["output_attivato_totale"]) + "**",
                     "**" + n2(s["moltiplicatore_medio"]) + "**",
                     "**" + n1(s["valore_aggiunto_totale"]) + "**",
                     "**" + n0(s["occupazione_ula_totale"]) + "**"])
    table(doc, ["Regione", "Domanda diretta (M€)", "Output totale (M€)", "Moltiplic.",
                "Valore aggiunto (M€)", "Occupazione (ULA)"], reg_rows, right_from=1, fs=8.5)
    caption(doc, "Tabella 2 — Sintesi della simulazione (valori rappresentativi, per anno a regime).")
    h2(doc, "3.4 Verifica del KPI")
    table(doc, ["Componente del deliverable", "Baseline", "Obiettivo", "Esito"],
          [["Matrice / modello input-output (Leontief) impostato e validato", "Nessuna informazione", "Analisi effettuata", "✓ prodotto"],
           ["Analisi eseguita in simulazione sulle 4 regioni", "Nessuna informazione", "Analisi effettuata", "✓ eseguita"],
           ["Protocollo di applicazione ex-post definito", "Nessuna informazione", "Analisi effettuata", "✓ definito"]],
          right_from=1, fs=8.5)
    caption(doc, "Tabella 3 — Verifica del KPI di scheda (baseline Nessuna informazione → obiettivo Analisi effettuata).")
    para(doc, "Il KPI **«Analisi effettuata»** è raggiunto nel senso proprio dell'attività: lo strumento "
              "di analisi input-output (matrice I-O) è **impostato, validato ed eseguito in simulazione** "
              "sulle quattro regioni, ed è corredato del protocollo per la sua applicazione ex-post.")

    # ============= 4. DISCUSSIONE E CONCLUSIONI =============
    h1(doc, "4. DISCUSSIONE E CONCLUSIONI")
    h2(doc, "4.1 Discussione critica")
    para(doc, f"La simulazione mostra che l'impatto potenziale del progetto **eccederebbe sensibilmente "
              f"lo shock diretto**: ogni euro di domanda finale attiva ~{n2(s['moltiplicatore_medio'])} € "
              "di output complessivo, con un contributo rilevante degli effetti indiretti (filiera) e "
              "indotti (consumi). La differenza tra le regioni non è un artefatto ma il riflesso della "
              "**struttura della filiera**: l'Emilia-Romagna, sede del core ceramico, presenta il "
              "moltiplicatore più elevato perché il settore ceramico attiva a monte estrazione, chimica, "
              "energia e logistica locali; le tre regioni della ricerca attivano soprattutto servizi e "
              "R&S, con legami a monte meno intensi ma comunque superiori all'unità. Vanno ribaditi i "
              "limiti che motivano la natura di simulazione: (i) i coefficienti sono rappresentativi e "
              "vanno sostituiti con le tavole I-O regionali ufficiali; (ii) l'analisi I-O assume "
              "coefficienti tecnici fissi e rendimenti costanti di scala, ipotesi ragionevole per impatti "
              "marginali ma non per grandi shock strutturali. È per questo che il deliverable è impostato "
              "come **modello da applicare ex-post**: i numeri odierni sono una previsione, il valore "
              "duraturo è lo strumento.")
    h2(doc, "4.2 Interdipendenze con altre attività e contributo a RF 8")
    bullet(doc, "**RP 8.5 (Stakeholder Engagement)** — è il **presupposto operativo** di questa analisi: "
                "il coinvolgimento delle autorità territoriali delle quattro regioni è la via di accesso "
                "ai dati macroeconomici regionali che risolvono il problema n. 6.")
    bullet(doc, "**Risultati tecnici (OR 1–OR 7)** — Digital Twin e AI (OR1), modelli predittivi di "
                "qualità e NDT (OR2), involucro edilizio intelligente (OR3), framework e controllo "
                "predittivo AI (OR4–OR5), modellazione dell'Intelligent Industry e product design (OR6), "
                "collaudo in ambiente operativo (OR7) definiscono la **sostanza economica** dello shock di "
                "domanda finale qui valutato.")
    bullet(doc, "**RF 8 (Piano di Coordinamento)** — la valutazione d'impatto è la componente di "
                "**misurazione dei risultati sui territori** del Piano di Coordinamento, e realizza la "
                "finalità inedita dell'OR 8.")
    h2(doc, "4.3 Implicazioni operative rispetto alle Finalità del progetto")
    para(doc, "Disporre di un modello per stimare l'impatto territoriale traduce in **numeri "
              "argomentabili** la Finalità di START di guidare la transizione dell'industria ceramica "
              "verso una produzione data-driven e sostenibile e di rafforzarne la competitività. Già in "
              "forma di simulazione i risultati forniscono elementi per il dialogo con le istituzioni "
              "(regioni, provincia autonoma, comuni, associazioni di distretto) e per la rendicontazione "
              "dell'Accordo per l'Innovazione, mostrando l'ordine di grandezza con cui gli effetti del "
              "progetto ricadrebbero sui territori di riferimento. L'applicazione ex-post trasformerà "
              "questa stima in evidenza consuntiva, dando all'OR 8 la sua «misurazione dei risultati sui "
              "territori».")
    h2(doc, "4.4 Limiti e sviluppi")
    bullet(doc, "**Applicazione ex-post (sviluppo principale).** Eseguire il protocollo del § 2.5: "
                "sostituire i coefficienti rappresentativi con le tavole I-O regionali ISTAT (via RP 8.5) "
                "e con i consuntivi degli investimenti.")
    bullet(doc, "**Analisi degli scostamenti.** Confrontare l'impatto simulato (questa relazione) con "
                "quello consuntivo ex-post, alimentando l'analisi degli scostamenti del Piano di "
                "Coordinamento.")
    bullet(doc, "**Estensione ambientale.** Affiancare un modulo input-output ambientale (emissioni, "
                "energia, exergia) agganciato alle impronte e alla contabilità EEA+/TSI del progetto (OR6).")
    bullet(doc, "**Robustezza della simulazione.** Corredare il modello di analisi di sensitività sui "
                "parametri rappresentativi, in attesa dei dati ufficiali.")
    h2(doc, "4.5 Conclusioni")
    para(doc, "L'attività 8.6 ha **impostato, validato ed eseguito in simulazione** il modello di analisi "
              "input-output dell'impatto del progetto START su Emilia-Romagna, Provincia di Bolzano, "
              "Sardegna e Calabria, e ne ha definito il **protocollo di applicazione ex-post** (KPI: "
              "baseline Nessuna informazione → obiettivo Analisi effettuata, raggiunto nel senso di "
              f"strumento predisposto ed eseguito). La simulazione stima un impatto potenziale di "
              f"**~{n1(s['output_attivato_totale'])} M€ di output**, **~{n1(s['valore_aggiunto_totale'])} "
              f"M€ di valore aggiunto** e **~{n0(s['occupazione_ula_totale'])} ULA** sui quattro territori "
              f"(moltiplicatore medio {n2(s['moltiplicatore_medio'])}), con l'Emilia-Romagna trainata dal "
              "core ceramico. Il valore del deliverable è però soprattutto il **modello riproducibile**, "
              "pronto a essere ri-eseguito con i dati regionali ufficiali — il cui accesso è abilitato dal "
              "Piano di Stakeholder Engagement (RP 8.5). Così l'attività contribuisce al **RF 8 (Piano di "
              "Coordinamento del Progetto START)** e mette l'OR 8 nelle condizioni di realizzare la sua "
              "finalità inedita: misurare l'impatto del progetto sulla società e sui territori di "
              "riferimento, per differenza tra previsione e consuntivo.")

    # ============= APPENDICI =============
    h1(doc, "Appendice A — Riproducibilità")
    para(doc, "Motore input-output (inversa di Leontief, moltiplicatori Tipo I/II, impatti): "
              "`RP8.6/src/impatto.py`. Runner: `RP8.6/run_impatto.py` (esporta "
              "`RP8.6/output/impatto_*.csv`). Figure: `RP8.6/scripts/gen_figures_impatto.py`. Questo "
              "documento è generato da `RP8.6/src/build_report.py` sul template ufficiale START.")
    h1(doc, "Appendice B — Riferimenti")
    bullet(doc, "Piano di Sviluppo START (Allegato 4, OR 8, attività 8.6); scheda KPI RP 8.6; problema "
                "progettuale n. 6; RP 8.5 (Piano di Stakeholder Engagement); risultati OR 1–OR 7. "
                "Impostazione metodologica gemella: progetto VOLT, RP 9.6.")
    bullet(doc, "Leontief, W. (1936). *Quantitative input and output relations in the economic systems of "
                "the United States.* The Review of Economics and Statistics, 18(3), 105–125.")
    bullet(doc, "Miller, R. E., & Blair, P. D. (2009). *Input-Output Analysis: Foundations and Extensions* "
                "(2nd ed.). Cambridge University Press.")
    bullet(doc, "Dietzenbacher, E., & Lahr, M. L. (Eds.) (2004). *Wassily Leontief and Input-Output "
                "Economics.* Cambridge University Press.")
    bullet(doc, "ISTAT. *Tavole delle risorse e degli impieghi (SUT) e tavole input-output* (fonte di "
                "consolidamento dei coefficienti regionali).")

    doc.save(OUTDOCX)
    print("DOCX generato:", OUTDOCX)


if __name__ == "__main__":
    build()
