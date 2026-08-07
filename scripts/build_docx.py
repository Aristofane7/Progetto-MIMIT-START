"""Genera il documento finale RP 7.10 in formato .docx rispettando la formattazione
del template ufficiale `RPX.Y Titolo_Relazione_Parziale_data.docx`.

Metodo:
  1. si parte dal reference-doc di default di pandoc e se ne ridefiniscono gli stili
     (Normal, Title, Heading 1/2/3, Table) per allinearli al template: corpo in
     Arial Narrow giustificato, titoli di sezione in grassetto, pagina A4 con i
     margini del template (L/R 1 cm, T/B 3 cm);
  2. pandoc converte il corpo della relazione (dalla sezione "## 1" in poi) usando
     quel reference-doc, gestendo tabelle, figure e formattazione inline;
  3. si antepone il blocco di testata (titolo + metadati) con la formattazione del
     template (titolo grande in grassetto centrato; metadati Calibri in grassetto
     centrato).

Uso: python scripts/build_docx.py
Requisiti: pandoc nel PATH, python-docx.
"""

from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
MD = DOCS / "RP7.10_AI_Business_Model.md"
TEMPLATE = ROOT / "RPX.Y Titolo_Relazione_Parziale_data.docx"
OUT = DOCS / "RP7.10_AI_Business_Model.docx"

# Formattazione allineata alle relazioni finali del progetto (es. RP 7.4):
#   corpo Arial Narrow 11 pt giustificato; titoli in Times New Roman grassetto
#   verde scuro (#024C41), 16 pt (sezioni) / 13 pt (sottosezioni); titolo di testata
#   Arial Narrow 28 pt grassetto; metadati Calibri 18 pt grassetto.
BODY_FONT = "Arial Narrow"        # font del corpo (Normal)
HEAD_FONT = "Times New Roman"     # font dei titoli di sezione
META_FONT = "Calibri"            # font dei metadati di testata
BLACK = RGBColor(0, 0, 0)
HEAD_COLOR = RGBColor(0x02, 0x4C, 0x41)   # verde scuro dei titoli

# Blocco di testata (come nelle RP finali): titolo + 4 righe di metadati.
TITLE = "Sviluppo di un modello di business basato sull'Intelligenza Artificiale (AI-BM)"
META_LINES = [
    "Relazione Parziale N°: RP7.10",
    "Versione del Documento: RV.1",
    "Data di Revisione del Documento: 07.08.2026",
    "Responsabilità: Gresmalt — Capofila",
]


def _style(doc, name):
    for s in doc.styles:
        if s.name == name:
            return s
    raise KeyError(name)


def _set_style(style, *, font=None, size=None, bold=None, italic=None, color=None):
    f = style.font
    if font is not None:
        f.name = font
        rpr = style.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rf.set(qn(a), font)
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = color


def build_reference() -> Path:
    """Crea un reference-doc pandoc con gli stili allineati al template."""
    ref = DOCS / ".ref_rp710.docx"
    subprocess.run(
        ["pandoc", "--print-default-data-file", "reference.docx"],
        check=True, stdout=ref.open("wb"),
    )
    d = Document(str(ref))

    # Corpo (Normal): Arial Narrow 11 pt, giustificato.
    normal = _style(d, "Normal")
    _set_style(normal, font=BODY_FONT, size=11, color=BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)

    # Stili del corpo usati da pandoc: basati su Normal (Arial Narrow 11).
    for nm in ("Body Text", "First Paragraph", "Compact"):
        try:
            _set_style(_style(d, nm), font=BODY_FONT, size=11, color=BLACK)
            _style(d, nm).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except KeyError:
            pass

    # Titoli: Times New Roman grassetto verde scuro (20/16/13 pt) come nelle RP finali.
    _set_style(_style(d, "Heading 1"), font=HEAD_FONT, size=20, bold=True, color=HEAD_COLOR)
    _set_style(_style(d, "Heading 2"), font=HEAD_FONT, size=16, bold=True, color=HEAD_COLOR)
    _set_style(_style(d, "Heading 3"), font=HEAD_FONT, size=13, bold=True, color=HEAD_COLOR)
    _style(d, "Heading 1").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for h in ("Heading 2", "Heading 3"):
        _style(d, h).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for h in ("Heading 1", "Heading 2", "Heading 3"):
        pf = _style(d, h).paragraph_format
        pf.space_before = Pt(12); pf.space_after = Pt(6)

    # Pagina A4 con i margini del template.
    sec = d.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.70)
    sec.left_margin = Cm(1.0); sec.right_margin = Cm(1.0)
    sec.top_margin = Cm(3.0); sec.bottom_margin = Cm(3.0)

    d.save(str(ref))
    return ref


def body_markdown() -> Path:
    """Estrae il corpo della relazione dalla sezione '## 1' in poi (testata esclusa)."""
    text = MD.read_text(encoding="utf-8")
    idx = text.index("## 1")
    body = DOCS / ".body_rp710.md"
    body.write_text(text[idx:], encoding="utf-8")
    return body


def run_pandoc(ref: Path, body: Path) -> None:
    # Nessuno shift: '## ' -> Heading 2 (sezioni), '### ' -> Heading 3 (sottosezioni).
    # '-implicit_figures': le immagini non generano una didascalia automatica dal testo
    # alternativo (si usa la sola riga di didascalia descrittiva della relazione).
    subprocess.run(
        ["pandoc", str(body.name), "--reference-doc", str(ref.name),
         "-f", "markdown-implicit_figures", "-o", OUT.name],
        check=True, cwd=str(DOCS),
    )


def _center_para(doc_para, text, *, font, size, bold=True, space_after=4):
    doc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_para.paragraph_format.space_after = Pt(space_after)
    r = doc_para.add_run(text)
    r.font.name = font          # python-docx imposta w:rFonts correttamente
    r.font.bold = bold
    r.font.size = Pt(size)


def prepend_header() -> None:
    """Antepone il blocco di testata (titolo + metadati) al documento finale."""
    d = Document(str(OUT))
    body_first = d.paragraphs[0]._element

    # Paragrafi creati nello stesso documento (evita problemi di namespace) e poi
    # spostati in testa, nell'ordine corretto.
    made = []
    p = d.add_paragraph(); _center_para(p, TITLE, font=BODY_FONT, size=28, space_after=10)
    made.append(p)
    for line in META_LINES:
        p = d.add_paragraph(); _center_para(p, line, font=META_FONT, size=18, space_after=3)
        made.append(p)

    # Riga separatrice: il bordo (w:pBdr) va inserito PRIMA di w:spacing nell'ordine
    # previsto da OOXML, quindi si aggiunge il bordo e solo dopo lo spazio inferiore.
    sep = d.add_paragraph()
    pPr = sep._element.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "000000")):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)
    sep.paragraph_format.space_after = Pt(10)
    made.append(sep)

    for p in made:
        body_first.addprevious(p._element)

    _style_captions_and_images(d)
    d.save(str(OUT))


def _style_captions_and_images(d) -> None:
    """Centra le immagini e porta le didascalie (Figura/Tabella) a 10 pt, come RP 7.4."""
    for p in d.paragraphs:
        if p._element.findall(".//" + qn("w:drawing")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        t = p.text.strip()
        if t.startswith("Figura ") or t.startswith("Tabella "):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)


def main() -> None:
    ref = build_reference()
    run_pandoc(ref, body_markdown())
    prepend_header()
    # pulizia file temporanei
    for f in (DOCS / ".ref_rp710.docx", DOCS / ".body_rp710.md"):
        f.unlink(missing_ok=True)
    print("Documento finale generato:", OUT)


if __name__ == "__main__":
    main()
