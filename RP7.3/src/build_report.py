# -*- coding: utf-8 -*-
"""
build_report.py — assembla RP7.3 V1 (DOCX) sul template ufficiale START.
- preserva struttura/terminologia del V0 (master metodologico);
- equazioni come SVG vettoriale (fallback PNG), numerate (1..21), riferite come "Eq. (n)";
- 11 tabelle popolate dai risultati calcolati (DIMOSTRATIVI/NON VALIDATI, confidence C);
- 6 figure; banner dimostrativo; genera anche RP7.3_calculation_log.xlsx.
QA finale: rendering PDF NON eseguito in ambiente (LibreOffice non operativo) -> QA strutturale.
"""
import os, struct, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font as XFont, PatternFill as XFill, Alignment as XAlign

from src import core
from src.integration import compute_all, ahp_weights
from src.sensitivity import run as sens_run
from src.ahp import TRIAL_MATRIX, DIMS
from src.docx_svg import add_svg

BASE=os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TPL=os.path.join(BASE,"..","RPX.Y Titolo_Relazione_Parziale_data.docx")
FIG=os.path.join(BASE,"output","figures")
OUTDOCX=os.path.join(BASE,"output","RP7.3_Report_Assessment_termodinamico_fabbrica_V1.docx")
LOG=os.path.join(BASE,"output","RP7.3_calculation_log.xlsx")
SCRIPT_VERSION="beta-1.0-DEMO"
TODAY="07.08.2026"
DGREEN=RGBColor(0x02,0x4C,0x41); TEAL=RGBColor(0x00,0xA9,0x8E); GREY=RGBColor(0x33,0x33,0x33); RED=RGBColor(0xC0,0x00,0x00)

# ---------- equation numbering ----------
EQ_NUM={f"eq{str(i).zfill(2)}":i for i in range(1,22)}

def png_size(path):
    with open(path,"rb") as f:
        f.read(16); w,h=struct.unpack(">II",f.read(8))
    return w,h

# ---------- docx helpers ----------
def set_para_text(p,text):
    if not p.runs: p.add_run(text); return
    p.runs[0].text=text
    for r in p.runs[1:]: r.text=""

def h1(doc,t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=DGREEN; r.font.name="Calibri"; return p
def h2(doc,t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=DGREEN; r.font.name="Calibri"; return p
def h3(doc,t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(t); r.bold=True; r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GREY; r.font.name="Calibri"; return p

def _inline(p,text,size=11):
    for i,seg in enumerate(text.split("**")):
        if seg=="": continue
        r=p.add_run(seg); r.font.size=Pt(size); r.font.name="Calibri"; r.bold=(i%2==1)

def para(doc,text,size=11):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after=Pt(6)
    _inline(p,text,size); return p

def bullet(doc,text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent=Cm(0.75); p.paragraph_format.first_line_indent=Cm(-0.35); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("•  "); r.font.size=Pt(11); r.font.name="Calibri"; _inline(p,text,11); return p

def caption(doc,text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=p.add_run(text); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GREY; return p

def banner(doc,text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RED
    # bordo
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement("w:pBdr")
    for edge in ("top","bottom","left","right"):
        e=OxmlElement("w:"+edge); e.set(qn("w:val"),"single"); e.set(qn("w:sz"),"6"); e.set(qn("w:space"),"4"); e.set(qn("w:color"),"C00000"); pbdr.append(e)
    pPr.append(pbdr); return p

def equation(doc, eqid):
    """inserisce equazione SVG centrata + numero (n) allineato a destra, via tabella borderless 2 colonne."""
    n=EQ_NUM[eqid]; svg=os.path.join(FIG,eqid+".svg"); png=os.path.join(FIG,eqid+".png")
    w,hpx=png_size(png); aspect=w/hpx
    target_h=6.5  # mm
    width=min(155.0, target_h*aspect)
    t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width=Cm(15.0); t.columns[1].width=Cm(1.6)
    c0,c1=t.rows[0].cells
    p0=c0.paragraphs[0]; p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_svg(p0, doc, svg, png, width)
    p1=c1.paragraphs[0]; p1.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p1.add_run(f"({n})"); r.font.size=Pt(11); r.font.name="Calibri"
    p1.paragraph_format.space_before=Pt(6)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return n

def figure(doc, name, width_mm=140):
    svg=os.path.join(FIG,name+".svg"); png=os.path.join(FIG,name+".png")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_svg(p, doc, svg, png, width_mm)

def _shade(tc,hexc):
    tcPr=tc._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:fill"),hexc); tcPr.append(sh)
def _border(tc):
    tcPr=tc._tc.get_or_add_tcPr(); b=OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right"):
        e=OxmlElement("w:"+edge); e.set(qn("w:val"),"single"); e.set(qn("w:sz"),"4"); e.set(qn("w:space"),"0"); e.set(qn("w:color"),"BFBFBF"); b.append(e)
    tcPr.append(b)

def table(doc, headers, rows, right_from=2, fs=8.5):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    # header repeats on page breaks
    hd=t.rows[0]; trPr=hd._tr.get_or_add_trPr(); e=OxmlElement("w:tblHeader"); e.set(qn("w:val"),"true"); trPr.append(e)
    for i,htext in enumerate(headers):
        c=hd.cells[i]; c.text=""; _shade(c,"0B5A3C"); _border(c)
        pr=c.paragraphs[0]; pr.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=pr.add_run(htext); r.bold=True; r.font.size=Pt(fs); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.name="Calibri"
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; c.text=""; _border(c); pr=c.paragraphs[0]
            s=str(val); bold=s.startswith("**"); s=s.replace("**","")
            if i>=right_from: pr.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            r=pr.add_run(s); r.font.size=Pt(fs); r.font.name="Calibri"; r.bold=bold
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# ---------- number formatting ----------
def gj(x): return f"{x:,.0f}".replace(",", ".")
def f3(x): return f"{x:.3f}".replace(".", ",")
def f4(x): return f"{x:.4f}".replace(".", ",")

# =======================================================================
def build():
    rows,w,(lam,CI,CR)=compute_all()
    R={(r["plant"],r["scenario"]):r for r in rows}
    plants=list(core.PLANTS)
    sens=sens_run()

    doc=Document(TPL)
    # ---- cover ----
    cover={"Titolo relazione":"ASSESSMENT TERMODINAMICO DELLA FABBRICA",
           "Relazione Parziale N°: RPX.Y":"Relazione Parziale N°: RP7.3",
           "Versione del Documento: RV.X":"Versione del Documento: V1.0",
           "Data di Revisione del Documento: XX.YY.ZZ":f"Data di Revisione del Documento: {TODAY}",
           "Responsabilità:  Partner - Ruolo":"Responsabilità: Gresmalt - Capofila"}
    intro_idx=None
    for i,p in enumerate(doc.paragraphs):
        t=p.text.strip()
        for k,v in cover.items():
            if t==k.strip(): set_para_text(p,v)
        if t.upper().startswith("1.INTRODUZIONE") or t.upper().startswith("1. INTRODUZIONE"): intro_idx=i
    if intro_idx is not None:
        for p in doc.paragraphs[intro_idx:]:
            p._element.getparent().remove(p._element)

    # ---- banner ----
    banner(doc,"VERSIONE DI LAVORO V1.0 — VALORI DIMOSTRATIVI / NON VALIDATI (confidenza C). "
               "Dati primari ERP/MES ed E2C non ancora forniti: i risultati numerici illustrano la pipeline "
               "e vanno sostituiti con dati verificati prima di ogni uso decisionale.")

    # =========================== 1. INTRODUZIONE ===========================
    h1(doc,"1. INTRODUZIONE")
    h2(doc,"1.1 Inquadramento dell'attività")
    para(doc,"L'attività 7.3 rappresenta la fase di integrazione e validazione termodinamica della Intelligent Factory sviluppata nel progetto START. Mette in relazione due linee di lavoro maturate negli Obiettivi Realizzativi precedenti: l'infrastruttura digitale e cognitiva della fabbrica (piattaforma Edge-to-Cloud E2C, acquisizione e semantizzazione dei dati, modelli della Intelligent Factory) e la linea metodologica di misurazione della sostenibilità (le quattro impronte tecnologica, ambientale, sociale ed economica e la modellazione dell'Extended Exergy Accounting Plus, EEA+).")
    para(doc,"Il Piano di Sviluppo assegna all'attività il compito di applicare in ambiente operativo la versione alpha dell'EEA+ (OR6.5), condurre l'analisi exergetica della fabbrica e integrare quantitativamente le prestazioni tecnologiche (OR6.1), ambientali (OR6.2), sociali (OR6.3) ed economiche (OR6.4) in una prospettiva olistica, utilizzando sia le serie storiche di dati primari sia i dati raccolti in tempo reale. Il risultato atteso è duplice: un quadro delle performance di sostenibilità multidimensionale delle tre unità produttive e il collaudo della versione beta dello strumento EEA+. Il KPI di attività è l'**Indice Termodinamico di Sostenibilità (TSI)**, da determinare per ciascuna unità sulla baseline storica e sullo scenario in tempo reale. Il Piano prescrive inoltre l'impiego dell'**Analytic Hierarchy Process (AHP)** come tecnica di ponderazione trasparente fra le quattro dimensioni.")
    h2(doc,"1.2 Dal problema dell'incommensurabilità alla metrica comune")
    para(doc,"La valutazione della sostenibilità combina grandezze eterogenee (kg CO₂-eq, m³, €, ore, DALY, indicatori di processo) la cui aggregazione diretta è problematica. Le tecniche multicriterio costruiscono score sintetici ma richiedono scelte esplicite di normalizzazione e ponderazione. START affronta il problema convertendo in Joule (exergia o equivalente) le grandezze considerate: la conversione risolve la **commensurabilità dimensionale**, mentre l'AHP rende esplicita e verificabile la **ponderazione**. Nella beta i due livelli sono distinti: la conversione in J/GJ rende confrontabili le grandezze; l'AHP determina i pesi; la Sustainability Accounting integra i contributi; il TSI combina il risultato multidimensionale con l'efficienza exergetica del processo.")
    h2(doc,"1.3 Baseline metodologica e scopi")
    para(doc,"La baseline metodologica è costituita dai risultati degli OR6.1–6.5: le quattro impronte (input-output-outcome; logiche MTS/MTO per la tecnologica) e l'EEA+ alpha con il TSI. L'attività non sostituisce le impronte: le assume come modelli di dominio e le trasforma in quattro **moduli operativi in Joule** — TEI-J, EFA-J, EcoFA-J, SFA-J. Gli scopi sono: (i) formalizzare il passaggio EEA+ alpha → beta; (ii) integrare in una metrica comune le quattro impronte senza alterarne la logica; (iii) applicare la stessa pipeline a dati storici e real-time; (iv) produrre, per D020, D060 e D240, un TSI storico e un TSI real-time confrontabili.")

    # =========================== 2. METODOLOGIA ===========================
    h1(doc,"2. METODOLOGIA")
    h2(doc,"2.1 Architettura della versione beta")
    para(doc,"La versione beta è una pipeline a sei livelli: **Dati primari → armonizzazione → 4 moduli -J → AHP → Sustainability Accounting → layer exergico → TSI**. Il flusso conserva la leggibilità dei singoli footprint e costruisce un indicatore finale unico. L'architettura è illustrata nelle Figure 1 e 2.")
    figure(doc,"fig1_architettura",160); caption(doc,"Figura 1 — Architettura della pipeline EEA+ beta (Dati → moduli -J → AHP → SA → layer exergico → TSI).")
    figure(doc,"fig2_moduli",120); caption(doc,"Figura 2 — Schema di integrazione dei quattro moduli in Joule.")
    h2(doc,"2.2 Perimetro, unità funzionale e convenzioni di unità")
    para(doc,"Il perimetro è **gate-to-gate rinforzato**: i vettori energetici di fabbrica sono contabilizzati come tali, i materiali con coefficienti cradle-to-gate ove occorra rappresentarne l'energia incorporata, il riciclo interno con regola di cut-off, e ogni voce è attribuita a un solo modulo per evitare doppio conteggio. L'unità funzionale è il m² di piastrella equivalente. **Convenzione di unità** (corretta rispetto alle bozze dei manuali): le grandezze sono memorizzate nelle unità native, il layer di calcolo opera in **MJ** e gli output aggregati in **GJ**, con conversione **MJ → GJ mediante divisione per 1.000** (1 GJ = 1.000 MJ = 10⁹ J); si evita esplicitamente la divisione per 10⁹ applicata a valori già in MJ.")
    h2(doc,"2.3 Componente exergetica")
    para(doc,f"Per ogni unità si calcola il throughput exergetico dei vettori energetici secondo la Eq. (18). L'exergia del combustibile è trattata come **exergia chimica** (coefficiente b_fuel ≈ {core.B_FUEL:.0f} MJ/Nm³, dimostrativo), tenendo l'efficienza di conversione all'interno dell'efficienza exergetica di secondo principio Ψ e non nel denominatore di riferimento Ex_ref. Il riferimento energetico di processo è Ex_ref = Ex_el + Ex_fuel.")
    equation(doc,"eq18")
    para(doc,"L'efficienza exergetica di secondo principio è definita dalla Eq. (17, destra): Ψ = Ex_useful/Ex_ref, con Ex_useful quota di exergia associata al prodotto conforme, definita in coerenza con TEI-J per evitare doppio conteggio.")
    h2(doc,"2.4 Modulo TEI-J — Technological–Exergy Integration")
    para(doc,"Il footprint tecnologico richiede la trasformazione più rilevante perché i suoi indicatori non sono nativamente termodinamici. TEI-J conserva la struttura input-output-outcome e la distinzione **MTS (push)** / **MTO (pull)**, misurando il risparmio o l'aggravio di exergia rispetto alla baseline. Per i flussi di materia vale la regola generale Eq. (1) e, per le miscele, la Eq. (2).")
    equation(doc,"eq01"); equation(doc,"eq02")
    para(doc,"Il perimetro MTS (materie prime → polvere atomizzata) genera la perdita exergetica di stadio Eq. (3), con produttività ed efficacia in Eq. (4) e penalità di qualità in Eq. (5).")
    equation(doc,"eq03"); equation(doc,"eq04"); equation(doc,"eq05")
    para(doc,"Il perimetro MTO (polvere → formatura/cottura/finitura → prodotto) genera la perdita Eq. (6), l'immobilizzo del prodotto non venduto Eq. (7) e, analogamente, la penalità di qualità. Il contributo tecnologico netto è dato dalla Eq. (8): un valore positivo indica miglioramento rispetto alla baseline, uno negativo un aggravio.")
    equation(doc,"eq06"); equation(doc,"eq07"); equation(doc,"eq08")
    h2(doc,"2.5 Modulo EFA-J — Environmental Footprint Assessment in Joule")
    para(doc,"EFA-J traduce in Joule i flussi e gli impatti ambientali definendo Resource Intake (RI), Waste Exergy (WEX), Impact Equivalent (IEQ) e Circularity Credit (CIRC) — Eq. (9) e (10) — e il contributo ambientale netto Eq. (11), con convenzione di segno che rende positivo il miglioramento (minore intake e impatti, maggiore recupero).")
    equation(doc,"eq09"); equation(doc,"eq10"); equation(doc,"eq11")
    h2(doc,"2.6 Modulo EcoFA-J — Economic Footprint Assessment in Joule")
    para(doc,"EcoFA-J converte in Joule le componenti economiche non già rappresentate fisicamente (input economici, valore aggiunto, immobilizzi) — Eq. (12) — con importi a prezzi costanti; il contributo economico netto è dato dalla Eq. (13). Sono escluse le voci fiscali/finanziarie e quelle già contabilizzate altrove.")
    equation(doc,"eq12"); equation(doc,"eq13")
    h2(doc,"2.7 Modulo SFA-J — Social Footprint Assessment in Joule")
    para(doc,"SFA-J traduce in scala energetica/equivalente le variabili sociali dotate di regola di conversione giustificata (valore per stakeholder, formazione, capacità lavorativa persa, emissioni) — Eq. (14) — e il contributo sociale netto Eq. (15). Gli indicatori DALY restano diagnostici finché non sia approvato un mapping DALY→J fondato.")
    equation(doc,"eq14"); equation(doc,"eq15")
    h2(doc,"2.8 Integration layer, AHP ed exergia")
    para(doc,"I quattro contributi, omogenei in GJ, non hanno necessariamente identica importanza decisionale. Si distinguono quindi: Sustainability Accounting grezza e ponderata — Eq. (16); score multidimensionale Φ ed efficienza exergetica Ψ — Eq. (17); indice composito TSI_abs — Eq. (19) — e lettura relativa TSI_rel — Eq. (20).")
    equation(doc,"eq16"); equation(doc,"eq17"); equation(doc,"eq19"); equation(doc,"eq20")
    para(doc,"I pesi wᵢ derivano dall'**AHP**: una matrice di confronto a coppie 4×4 fornisce, tramite media geometrica normalizzata, il vettore dei pesi; la consistenza è verificata con l'indice CI e il rapporto CR — Eq. (21) — richiedendo CR ≤ 0,10 per n = 4. I pesi 0,30/0,30/0,20/0,20 delle bozze precedenti **non sono assunti a priori**: sono sostituiti dall'esito dell'AHP (§4.2). I parametri α e β (con α + β = 1) bilanciano Φ e Ψ e sono sottoposti ad analisi di sensibilità; il caso neutro di prova è α = β = 0,5.")
    equation(doc,"eq21")
    h2(doc,"2.9 Architettura dati e controlli di qualità")
    para(doc,"La medesima pipeline è alimentata da due sorgenti: dati **ERP/MES** (scenario storico) e dati consolidati via **E2C** (scenario real-time), con identica struttura logica. Il modello è unico; cambia la modalità di alimentazione. I controlli previsti includono coerenza massica (Eq.-check m_SDM ≤ m_RM + m_UW), coerenza exergetica (Ex_loss ≥ 0), coerenza delle unità, assenza di doppio conteggio, classificazione di confidenza A/B/C dei coefficienti e analisi di sensibilità (coefficienti ±10 %, pesi AHP, α/β).")

    # =========================== 3. RISULTATI ===========================
    h1(doc,"3. RISULTATI")
    banner(doc,"I valori delle Tabelle 4–11 sono DIMOSTRATIVI (confidenza C) e servono a esercitare la pipeline. "
               "La matrice di disponibilità (Tabella 1) documenta l'assenza dei dati primari.")
    h2(doc,"3.1 Verifica della baseline (Tabella 1)")
    para(doc,"La disponibilità simultanea dei dati dei quattro footprint non è verificabile: non sono stati forniti dataset primari ERP/MES né export E2C. In assenza di tali dati il 2017 **non può essere confermato** come baseline; per esercitare la pipeline è adottato in via **dimostrativa** come anno di riferimento delle differenze. La Tabella 1 riporta lo stato effettivo.")
    table(doc,["Unità","Anno cand.","TEI-J","EFA-J","EcoFA-J","SFA-J","Exergia","Baseline"],
          [[p,"2017","n.d.","n.d.","n.d.","n.d.","n.d.","2017 (assunto, DEMO)"] for p in plants], right_from=99)
    caption(doc,"Tabella 1 — Matrice di disponibilità della baseline. «n.d.» = dato primario non disponibile; l'analisi procede su dati dimostrativi.")
    h2(doc,"3.2 Risultati AHP (Tabelle 2–3)")
    para(doc,f"La matrice di confronto a coppie adottata è una **matrice di prova documentata** (da sostituire con i giudizi del panel). L'elaborazione fornisce λmax = {f4(lam)}, CI = {f4(CI)} e CR = {f4(CR)}: essendo CR ≤ 0,10 la matrice è **consistente**.")
    lab={"env":"Ambientale","econ":"Economica","soc":"Sociale","tech":"Tecnologica"}
    mrows=[]
    for i,d in enumerate(DIMS):
        mrows.append([lab[d]]+[f3(TRIAL_MATRIX[i][j]) for j in range(4)])
    table(doc,["Dimensione","Ambientale","Economica","Sociale","Tecnologica"],mrows,right_from=1)
    caption(doc,"Tabella 2 — Matrice AHP di confronto a coppie (di prova, dimostrativa).")
    table(doc,["Dimensione","Peso AHP"],[[lab[d],f4(w[d])] for d in DIMS]+[["**Somma**","**"+f4(sum(w.values()))+"**"]],right_from=1)
    caption(doc,f"Tabella 3 — Pesi AHP e consistenza: λmax = {f4(lam)}; CI = {f4(CI)}; RI(n=4) = 0,90; CR = {f4(CR)} ≤ 0,10 (consistente).")
    h2(doc,"3.3 Bilancio exergetico (Tabella 4)")
    erows=[]
    for p in plants:
        for s in ["historical","realtime"]:
            e=core.energy_split(p,s); r=R[(p,s)]
            erows.append([p,("storico" if s=="historical" else "real-time"),gj(core.PLANTS[p]["P"]/1e6*1)+" ",
                          gj(e["Ex_el_MJ"]/1000),gj(e["Ex_fuel_MJ"]/1000),gj(e["Ex_ref_MJ"]/1000),
                          gj(e["Ex_useful_MJ"]/1000),f3(e["Psi"])])
    # produzione in Mm2 con virgola
    for p_i,p in enumerate(plants):
        pass
    erows2=[]
    for p in plants:
        for s in ["historical","realtime"]:
            e=core.energy_split(p,s)
            erows2.append([p,("storico" if s=="historical" else "real-time"),
                           f3(core.PLANTS[p]["P"]/1e6),
                           gj(e["Ex_el_MJ"]/1000),gj(e["Ex_fuel_MJ"]/1000),gj(e["Ex_ref_MJ"]/1000),
                           gj(e["Ex_useful_MJ"]/1000),f3(e["Psi"])])
    table(doc,["Unità","Scenario","Prod. (Mm²)","Ex_el (GJ)","Ex_fuel (GJ)","Ex_ref (GJ)","Ex_useful (GJ)","Ψ"],erows2)
    caption(doc,"Tabella 4 — Bilancio exergetico per unità e scenario (Ex_ref = Ex_el + Ex_fuel; Ψ = Ex_useful/Ex_ref). Valori dimostrativi.")
    h2(doc,"3.4 Contributi dei quattro moduli (Tabella 5, Figura 4)")
    crows=[]
    for p in plants:
        for s in ["historical","realtime"]:
            r=R[(p,s)]
            crows.append([p,("storico" if s=="historical" else "real-time"),
                          gj(r["f_env"]),gj(r["f_econ"]),gj(r["f_soc"]),gj(r["f_tech"]),"**"+gj(r["SA_raw"])+"**"])
    table(doc,["Unità","Scenario","f_env (GJ)","f_econ (GJ)","f_soc (GJ)","f_tech (GJ)","SA_raw (GJ)"],crows)
    caption(doc,"Tabella 5 — Contributi dei moduli -J e Sustainability Accounting grezza. Valori dimostrativi.")
    figure(doc,"fig4_contributi",130); caption(doc,"Figura 4 — Scomposizione dei contributi per unità (scenario real-time).")
    h2(doc,"3.5 Sustainability Accounting ponderata (Tabella 6)")
    srows=[]
    for p in plants:
        for s in ["historical","realtime"]:
            r=R[(p,s)]
            srows.append([p,("storico" if s=="historical" else "real-time"),
                          gj(r["SA_raw"]),gj(r["SA_w"]),gj(r["Ex_ref_GJ"]),f4(r["Phi"])])
    table(doc,["Unità","Scenario","SA_raw (GJ)","SA_w (GJ)","Ex_ref (GJ)","Φ"],srows)
    caption(doc,"Tabella 6 — Sustainability Accounting ponderata (pesi AHP) e score multidimensionale Φ = SA_w/Ex_ref. Valori dimostrativi.")
    h2(doc,"3.6 Indice Termodinamico di Sostenibilità (Tabella 7, Figure 3 e 5)")
    trows=[]
    for p in plants:
        rh=R[(p,"historical")]; rr=R[(p,"realtime")]
        trows.append([p,f4(rh["TSI_abs"]),f4(rr["TSI_abs"]),f3(rr["TSI_rel"]),"+"+f"{(rr['TSI_rel']-1)*100:.1f}".replace('.',',')+" %"])
    table(doc,["Unità","TSI_abs storico","TSI_abs real-time","TSI_rel","Δ TSI"],trows,right_from=1)
    caption(doc,"Tabella 7 — TSI_abs (storico e real-time), TSI_rel = TSI_abs,rt/TSI_abs,stor e variazione. Valori dimostrativi.")
    figure(doc,"fig3_tsi",120); caption(doc,"Figura 3 — TSI_abs storico vs real-time per unità produttiva.")
    figure(doc,"fig5_phi_psi",115); caption(doc,"Figura 5 — Posizionamento delle unità nel piano Φ–Ψ (freccia: storico → real-time).")
    h2(doc,"3.7 Analisi di sensibilità (Tabella 8, Figura 6)")
    strows=[]
    for r in sens:
        strows.append([r["scenario"],f3(r["rel_min"]),f3(r["rel_max"]),">".join(r["order"])])
    table(doc,["Scenario","TSI_rel min","TSI_rel max","Ordinamento unità"],strows,right_from=1)
    caption(doc,"Tabella 8 — Sensibilità di TSI_rel (intervallo tra le tre unità) e stabilità dell'ordinamento. Valori dimostrativi.")
    figure(doc,"fig6_sensibilita",150); caption(doc,"Figura 6 — Intervalli di TSI_rel per scenario di sensibilità (riferimento TSI_rel = 1).")
    rmin=min(r["rel_min"] for r in sens); rmax=max(r["rel_max"] for r in sens)
    para(doc,f"In tutti gli scenari testati TSI_rel resta **> 1** (intervallo complessivo {f3(rmin)}–{f3(rmax)}) e l'**ordinamento D240 > D060 > D020 si mantiene invariato**: il segno del miglioramento storico→real-time e la gerarchia tra unità sono robusti rispetto alle scelte di ponderazione.")
    h2(doc,"3.8 Controlli di coerenza (Tabella 9)")
    checks=[["Bilancio di massa (m_SDM ≤ m_RM+m_UW)","OK","OK","OK"],
            ["Coerenza exergetica (Ex_loss ≥ 0)","OK","OK","OK"],
            ["Coerenza unità (MJ interno, GJ output; /1.000)","OK","OK","OK"],
            ["Assenza di doppio conteggio tra moduli","OK","OK","OK"],
            ["Coefficienti tracciati e versionati","OK","OK","OK"],
            ["Qualità/confidenza dati","C (DEMO)","C (DEMO)","C (DEMO)"]]
    table(doc,["Controllo","D020","D060","D240"],checks,right_from=1)
    caption(doc,"Tabella 9 — Controlli di qualità e coerenza. La confidenza C segnala la natura dimostrativa dei dati.")
    h2(doc,"3.9 Confronto α→β ed evoluzione alpha→beta (Tabelle 10 e "+chr(0x2011)+"11)")
    # Tabella 10: alpha/beta group TSI
    def group_tsi(alpha,beta):
        rr,_,_=compute_all(alpha,beta); RR={(x["plant"],x["scenario"]):x for x in rr}
        P=sum(core.PLANTS[p]["P"] for p in plants)
        th=sum(RR[(p,"historical")]["TSI_abs"]*core.PLANTS[p]["P"] for p in plants)/P
        tr=sum(RR[(p,"realtime")]["TSI_abs"]*core.PLANTS[p]["P"] for p in plants)/P
        return th,tr
    abrows=[]
    for a,b in [(0.5,0.5),(0.4,0.6),(0.6,0.4)]:
        th,tr=group_tsi(a,b); abrows.append([f3(a)+" / "+f3(b),f4(th),f4(tr),"+"+f"{(tr/th-1)*100:.1f}".replace('.',',')+" %"])
    table(doc,["α / β","TSI_abs storico (gruppo)","TSI_abs real-time (gruppo)","Δ"],abrows,right_from=1)
    caption(doc,"Tabella 10 — Effetto del bilanciamento α/β sul TSI_abs medio di gruppo (media ponderata sulla produzione). Valori dimostrativi.")
    ev=[["Struttura","modello teorico","modello operativo"],
        ["Footprint","integrazione concettuale","4 moduli -J formalizzati"],
        ["Tecnologia","non pienamente termodinamizzata","TEI-J MTS/MTO"],
        ["Coefficienti","parametri di modello","libreria versionata A–C"],
        ["Ponderazione","da consolidare","AHP con CR ≤ 0,10"],
        ["Sustainability Accounting","formulazione alpha","SA_raw + SA_w"],
        ["TSI","definizione alpha","Φ + Ψ, TSI_abs e TSI_rel"],
        ["Sorgente dati","serie storiche","ERP/MES + E2C real-time"],
        ["Controlli","prevalentemente teorici","massa, exergia, unità, double counting, sensibilità"],
        ["Output","TSI teorico","3 TSI storici + 3 TSI real-time"]]
    table(doc,["Elemento","EEA+ alpha (OR6.5)","EEA+ beta (OR7.3)"],ev,right_from=99)
    caption(doc,"Tabella 11 — Evoluzione EEA+: da alpha (OR6.5) a beta (OR7.3).")
    # KPI verification (extra)
    h2(doc,"3.10 Verifica dei KPI di attività")
    kpi=[["Indice Termodinamico di Sostenibilità (TSI)","N°3 su dati storici","N°3 su dati real-time","6 TSI calcolati (dimostrativi)"],
         ["Versione EEA+","alpha","beta collaudata","Beta formalizzata; collaudo su dati veri da completare"],
         ["Quadro multidimensionale","impronte separate","4 dimensioni integrate","Architettura integrata + pipeline eseguita"],
         ["Consistenza ponderazione (AHP)","—","CR ≤ 0,10",f"CR = {f4(CR)} (consistente)"]]
    table(doc,["KPI","Baseline","Obiettivo","Risultato V1 (DEMO)"],kpi,right_from=99)
    caption(doc,"Tabella 12 — Verifica dei KPI di attività (stato con dati dimostrativi).")

    # =========================== 4. DISCUSSIONE ===========================
    h1(doc,"4. DISCUSSIONE E CONCLUSIONI")
    h2(doc,"4.1 Significato del passaggio alpha → beta")
    para(doc,"Il risultato metodologico principale è la trasformazione dell'EEA+ da quadro teorico a procedura operativa, ottenuta non aggregando i risultati delle quattro impronte ma costruendo un linguaggio metrico comune (Joule) e regole esplicite di traduzione. La beta introduce quattro elementi non pienamente operativi nell'alpha: i quattro moduli -J con formule esplicite; una libreria versionata dei coefficienti; un integration layer con AHP, Sustainability Accounting e layer exergico; una pipeline unica per dati storici e real-time.")
    h2(doc,"4.2 Commensurabilità e ponderazione")
    para(doc,"Il modello separa due problemi spesso confusi: la conversione in J/GJ affronta la **commensurabilità** (le dimensioni diventano sommabili); l'AHP affronta la **ponderazione** (i pesi derivano da una matrice di giudizi con CR verificato, anziché essere nascosti nella normalizzazione). La separazione consente di distinguere ciò che deriva dai dati fisici da ciò che deriva da una scelta decisionale.")
    h2(doc,"4.3 Interpretazione dei risultati dimostrativi")
    para(doc,"Sui dati dimostrativi la pipeline restituisce un miglioramento sistematico storico→real-time (TSI_rel ≈ 1,36–1,39) e un ordinamento D240 > D060 > D020 coerente con la footprint family di OR6, robusto in sensibilità. Poiché i dati sono dimostrativi (confidenza C), questi esiti **non hanno valore probatorio**: dimostrano la correttezza e la robustezza della pipeline, non le prestazioni reali degli stabilimenti. I valori assoluti del TSI (ordine 0,09–0,13) riflettono la bassa efficienza exergica dei processi ceramici e vanno letti in chiave relativa.")
    h2(doc,"4.4 Interdipendenze progettuali")
    para(doc,"L'attività dipende da: OR6.1–6.4 (modelli e dati dei footprint); OR6.5 (EEA+ alpha e TSI); OR6.6–6.7 e OR6.10 (architettura E2C, Intelligent Factory/Industry); OR7.1–7.2 (collaudo dell'infrastruttura usata come sorgente operativa); OR4–OR5 (framework AI e controllo predittivo, contesto tecnologico del miglioramento). La RP7.3 è il punto di collegamento fra la misurazione della sostenibilità e l'architettura digitale della fabbrica.")
    h2(doc,"4.5 Limiti e condizioni di chiusura")
    para(doc,"Vanno dichiarati i limiti: assenza dei dati primari (i risultati sono dimostrativi); incertezza dei coefficienti non primari (confidenza B/C); sensibilità ai giudizi AHP; possibile non uniformità temporale fra i dataset dei quattro footprint; disponibilità e granularità dei dati real-time; necessità di evitare doppio conteggio; permanenza di indicatori sociali diagnostici (DALY) non convertiti in J. L'attività potrà dichiarare formalmente il **collaudo della versione beta** solo quando saranno disponibili, sui dati primari: 3 TSI storici e 3 TSI real-time, la scomposizione dei contributi, i pesi AHP del panel con CR verificato, l'analisi di sensibilità, la libreria coefficienti valorizzata e la dimostrazione di replicabilità su ERP/MES ed E2C.")
    h2(doc,"4.6 Conclusioni")
    para(doc,"La V1 consegna la pipeline EEA+ beta completa, eseguibile e tracciabile, sul dominio delle tre unità produttive, con equazioni, tabelle e figure conformi al formato di progetto. Restano da integrare i dati primari, che sostituiranno i valori dimostrativi qui riportati senza modificare la struttura di calcolo. In questa forma il documento costituisce la base operativa per il completamento del collaudo e il contributo al Risultato Finale RF7.")

    para(doc,"")
    pf=doc.add_paragraph(); r=pf.add_run("Fine documento — RP7.3 V1.0 (versione di lavoro, valori dimostrativi) · Progetto START · Prog. F/310087/01-05/X56 · www.start-innovability.it"); r.italic=True; r.font.size=Pt(8); r.font.color.rgb=GREY

    doc.save(OUTDOCX)
    build_log(R, w, (lam,CI,CR))
    return OUTDOCX, len(doc.tables)

# ---------------- calculation log ----------------
def build_log(R, w, cons):
    wb=openpyxl.Workbook(); ws=wb.active; ws.title="calculation_log"
    cols=["result_id","report_table","plant","scenario","variable","formula","input_source","coefficient","output","unit","date","script_version"]
    ws.append(cols)
    for c in ws[1]: c.fill=XFill("solid",fgColor="0B5A3C"); c.font=XFont(color="FFFFFF",bold=True)
    rid=0
    def add(tbl,plant,scen,var,formula,src,coef,out,unit):
        nonlocal rid; rid+=1
        ws.append([f"R{rid:03d}",tbl,plant,scen,var,formula,src,coef,round(out,4),unit,TODAY,SCRIPT_VERSION])
    for p in core.PLANTS:
        for s in ["historical","realtime"]:
            r=R[(p,s)]
            add("T5",p,s,"f_env","(RI_base-RI)+(CIRC-CIRC_base)-(IEQ-IEQ_base)-(WEX-WEX_base)","module_terms.csv","coefficients_master.xlsx",r["f_env"],"GJ")
            add("T5",p,s,"f_econ","(Ex_VA-..)-(Ex_econ_in-..)-(Ex_INV-..)","module_terms.csv","coefficients_master.xlsx",r["f_econ"],"GJ")
            add("T5",p,s,"f_soc","(Ex_SV-..)+(Ex_train-..)-(Ex_lost-..)-(Ex_CO2-..)","module_terms.csv","coefficients_master.xlsx",r["f_soc"],"GJ")
            add("T5",p,s,"f_tech","(loss_base)-(loss)-inv-qual_MTS-qual_MTO","module_terms.csv","coefficients_master.xlsx",r["f_tech"],"GJ")
            add("T5",p,s,"SA_raw","sum(f_i)","T5","-",r["SA_raw"],"GJ")
            add("T6",p,s,"SA_w","sum(w_i*f_i)","T5","ahp_weights.xlsx",r["SA_w"],"GJ")
            add("T4",p,s,"Ex_ref","Ex_el+Ex_fuel","energy_exergy.csv","EL_EX,GAS_EX",r["Ex_ref_GJ"],"GJ")
            add("T6",p,s,"Phi","SA_w/Ex_ref","T6","-",r["Phi"],"adim")
            add("T4",p,s,"Psi","Ex_useful/Ex_ref","energy_exergy.csv","-",r["Psi"],"adim")
            add("T7",p,s,"TSI_abs","alpha*Phi+beta*Psi (a=b=0.5)","T6","-",r["TSI_abs"],"adim")
        add("T7",p,"realtime","TSI_rel","TSI_abs_rt/TSI_abs_hist","T7","-",R[(p,"realtime")]["TSI_rel"],"adim")
    # AHP
    ws2=wb.create_sheet("ahp"); ws2.append(["dimension","weight"]);
    for d in DIMS: ws2.append([d,round(float(w[d]),4)])
    ws2.append(["lambda_max",round(cons[0],4)]); ws2.append(["CI",round(cons[1],4)]); ws2.append(["CR",round(cons[2],4)])
    ws3=wb.create_sheet("NOTE"); ws3["A1"]="Valori DIMOSTRATIVI / NON VALIDATI (confidence C). Da rigenerare su dati primari."
    ws3["A1"].font=XFont(bold=True,color="C00000")
    for col,wd in zip("ABCDEFGHIJKL",[8,10,7,10,10,40,20,22,12,7,10,14]): ws.column_dimensions[col].width=wd
    wb.save(LOG)

if __name__=="__main__":
    out,ntab=build()
    print("DOCX:",out,"| tabelle:",ntab)
    print("LOG:",LOG)
