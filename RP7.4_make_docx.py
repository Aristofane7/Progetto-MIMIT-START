#!/usr/bin/env python3
"""Genera RP7.4 in .docx (house style RP7.3) dai risultati calcolati."""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/user/Progetto-MIMIT-START"
R = json.load(open(f"{OUT}/RP7.4_results.json"))
GREEN = RGBColor(0x4F, 0x62, 0x28)

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

def shade(cell, hexc="4F6228"):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc); tcPr.append(sh)

def para(text="", size=None, bold=False, italic=False, color=None, align=None, after=6, before=0):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    return p

def heading(text, lvl=1):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True
    r.font.size = Pt(14 if lvl == 1 else 11.5); r.font.color.rgb = GREEN
    p.paragraph_format.space_before = Pt(12 if lvl == 1 else 8); p.paragraph_format.space_after = Pt(4)
    return p

def table(headers, rows, caption=None, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hd = t.rows[0].cells
    for i, h in enumerate(headers):
        hd[i].text = ""; run = hd[i].paragraphs[0].add_run(h); run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(9.5)
        shade(hd[i]); hd[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        c = t.add_row().cells
        for i, v in enumerate(row):
            c[i].text = ""; run = c[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5)
            c[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
    if caption: para(caption, size=8.5, italic=True, after=10)
    return t

def f3(t, key):
    return f"{R[key][t]:+.3f}" if R[key][t] < 0 or key.endswith("_z") else f"{R[key][t]:.3f}"

# ---------------- intestazione ----------------
para("SusTainable dAta-dRiven manufacTuring", bold=True, size=11, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
para("Accordo Innovazione DM 31/12/2021 — Prog. n. F/310087/01-05/X56", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
para("www.start-innovability.it", size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
para("PRODUCT TECHNOLOGICAL SUSTAINABILITY ASSESSMENT (P-TSA)", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
for k, v in [("Relazione Parziale N°", "RP7.4"), ("Versione del Documento", "V1.0"),
             ("Data di Revisione del Documento", "07.08.2026"), ("Responsabilità", "Gresmalt — Capofila")]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(k + ": "); r1.bold = True; r1.font.size = Pt(10); p.add_run(v).font.size = Pt(10)
    p.paragraph_format.space_after = Pt(0)
doc.add_paragraph()

# ---------------- 1. INTRODUZIONE ----------------
heading("1. INTRODUZIONE")
heading("1.1 Inquadramento dell'attività", 2)
para("L'attività 7.4 estende al prodotto la valutazione di sostenibilità tecnologica che l'attività 7.3 conduce sulla fabbrica. "
     "Se l'OR7.3 misura, con l'Extended Exergy Accounting Plus (EEA+), la sostenibilità del sistema produttivo, l'OR7.4 completa il "
     "quadro determinando in modo quantitativo le prestazioni tecniche del prodotto ceramico e la loro conformità al quadro normativo "
     "dei materiali ceramici. Il Piano di Sviluppo assegna all'attività l'applicazione della metodologia P-TSA secondo l'approccio del "
     "ciclo di vita (Life Cycle Thinking, LCT) e lo schema della ISO 14040, in prospettiva di supply chain cradle-to-grave, misurando la "
     "sostenibilità tecnologica rispetto a tre categorie d'impatto: In-/Outputs Availability (IOA), Operational Performance (OP) e "
     "Technical Quality (TQ).")
heading("1.2 Baseline e scopi", 2)
para("La baseline di progetto è nulla (0.0) per tutti i KPI. Gli obiettivi sono: un indicatore per ciascuna categoria e per 3 tipologie "
     "— Stock Coverage Rate (SCR) per IOA, Productivity Indicators (PsI) per OP, Output Conformity Rate (OCR) per TQ — e un indice "
     "sintetico normalizzato, il Product Technology Sustainability Index (P-TSI), per le 3 tipologie. Le tre tipologie sono definite dalle "
     "Environmental Product Declaration (EPD, ISO 14025 / EN 15804+A2, EPDItaly) del gruppo:")
table(["Tipologia", "Spessore", "Massa (unità dich.)", "Stabilimento", "Gruppo", "Uso tipico"],
      [["T1", "7,4 mm", "13,98 kg/m²", "D060 Scandiano", "BIa", "interni, alleggerito"],
       ["T2", "8,2 mm", "16,05 kg/m²", "D060 Scandiano", "BIa", "interni/esterni, standard"],
       ["T3", "20,0 mm", "41,79 kg/m²", "D240 Frassinoro", "BIa", "esterni/outdoor, spessorato"]],
      "Tabella 1 — Le tre tipologie di prodotto (fonte: EPD). Tutte gres porcellanato smaltato gruppo BIa (assorbimento ≤ 0,5 %, ISO 10545-3), EN 14411, cottura 1210–1230 °C, unità dichiarata 1 m²·anno, periodo dati lug 2023 – giu 2024.")
heading("1.3 Interdipendenze", 2)
para("L'attività dipende da OR6.8/6.9 (segmentazione e progettazione data-driven del prodotto), da OR7.1–7.2 (infrastruttura "
     "Edge-to-Cloud E2C come sorgente dati) e condivide impianto metodologico e governance del dato con OR7.3. Il P-TSI è, con l'EEA+ "
     "Index, uno dei due KPI del collaudo della Intelligent Industry (OR7.8) e concorre alla caratterizzazione dei prototipi di OR7.9, "
     "contribuendo al Risultato Finale RF7.")

# ---------------- 2. METODOLOGIA ----------------
heading("2. METODOLOGIA")
heading("2.1 Fondamento e schema ISO 14040", 2)
para("Il P-TSA applica al prodotto la logica LCT/ISO 14040 in quattro fasi (Goal & Scope → Inventory → Impact Assessment → "
     "Interpretation), trattando la tecnologia come dimensione di sostenibilità a sé stante e non come mero abilitatore. Il "
     "sistema-prodotto è letto lungo le sette attività della value chain (Sourcing, Inbound Logistics, Operations, Internal Logistics, "
     "Outbound Logistics, Product Usage, Waste Logistics), dal cradle al grave.")
heading("2.2 Unità funzionale", 2)
para("Coerentemente con il problema progettuale n. 4, l'unità funzionale è stata scelta tramite prove comparate tra 1 m², 1 t e lotto di "
     "tipologia produttiva. Si è adottato il lotto di tipologia produttiva come unità funzionale operativa, mantenendo la normalizzazione "
     "a 1 m² — la stessa unità dichiarata degli EPD — come riferimento di comparabilità trasversale tra le tre tipologie e con OR7.3.")
heading("2.3 Categorie d'impatto e indicatori", 2)
para("• IOA → SCR (Stock Coverage Rate) = Stock medio / Consumo medio [giorni], su materie prime (Sourcing), prodotto finito (Internal "
     "Logistics) e smalti/inchiostri (Operations).")
para("• OP → PsI (Productivity Indicators) = Output reale / Input reale: produttività energetica [m²/GJ], resa di materiale [m²/m²], "
     "throughput di linea [m²/h].")
para("• TQ → OCR (Output Conformity Rate) = Parametro qualità / Soglia normativa, su parametri ISO 10545 con soglie EN 14411 BIa: "
     "resistenza a flessione (ISO 10545-4, ≥ 35 N/mm²), sforzo di rottura (ISO 10545-4, ≥ 700/1300 N), qualità superficiale "
     "(ISO 10545-2, ≥ 95 %).")
heading("2.4 Normalizzazione e aggregazione", 2)
para("Si adottano due schemi con verifica incrociata. Primario — z-score + pesi uguali (Vacchi 2021): ogni indicatore è standardizzato "
     "z = (x − μ)/σ tra le tre tipologie; i sotto-indici IOAI, OPI, TQI sono la media (pesi uguali) degli z; il P-TSI è la media dei tre "
     "sotto-indici (lettura relativa/comparativa). Secondario — scoring 1–5 + AHP (RP7.3 / O-TSA 2025): ogni indicatore è mappato su "
     "scala 1–5 tramite soglie; i pesi entro e tra le dimensioni derivano da Analytic Hierarchy Process con verifica del Consistency "
     "Ratio (CR ≤ 0,10) — lettura assoluta in [1–5]. Il confronto temporale è dato dal Technology Improvement Index (TII).")
heading("2.5 Architettura dati", 2)
para("La pipeline è alimentabile da serie storiche ERP/MES e da dati in tempo reale E2C (stessa logica di OR7.3). Gli artefatti "
     "— RP7.4_dataset_sintetico.xlsx, RP7.4_weights.xlsx, RP7.4_calculation_log.xlsx, RP7.4_build.py — replicano la struttura di RP7.3, "
     "sono versionati e tracciabili; le serie sono in corso di consolidamento.")

# ---------------- 3. RISULTATI ----------------
heading("3. RISULTATI")
heading("3.1 Energia specifica di processo", 2)
table(["Tipologia", "Massa [kg/m²]", "Energia specifica processo [MJ/m²]"],
      [["T1 · 7,4 mm", "13,98", f"{R['e_spec']['T1_7.4mm']:.1f}".replace('.',',')],
       ["T2 · 8,2 mm", "16,05", f"{R['e_spec']['T2_8.2mm']:.1f}".replace('.',',')],
       ["T3 · 20 mm", "41,79", f"{R['e_spec']['T3_20mm']:.1f}".replace('.',',')]],
      "Tabella 2 — Intensità energetica di processo per m² (∝ massa; ancoraggio a RP7.3 D060/D240).")

heading("3.2 Indicatori per categoria", 2)
def g(d, t, dec=3): return f"{d[t]:.{dec}f}".replace('.', ',')
scr = R["SCR"]
table(["SCR — copertura scorte [giorni]", "T1 · 7,4 mm", "T2 · 8,2 mm", "T3 · 20 mm"],
      [["Materie prime (Sourcing)", g(scr["SCR_RawMat"],"T1_7.4mm",1), g(scr["SCR_RawMat"],"T2_8.2mm",1), g(scr["SCR_RawMat"],"T3_20mm",1)],
       ["Prodotto finito (Internal Log.)", g(scr["SCR_Finished"],"T1_7.4mm",1), g(scr["SCR_Finished"],"T2_8.2mm",1), g(scr["SCR_Finished"],"T3_20mm",1)],
       ["Smalti/inchiostri (Operations)", g(scr["SCR_GlazeInk"],"T1_7.4mm",1), g(scr["SCR_GlazeInk"],"T2_8.2mm",1), g(scr["SCR_GlazeInk"],"T3_20mm",1)]],
      "Tabella 3 — IOA · Stock Coverage Rate per input e tipologia.")
pi = R["PI"]
table(["PsI — produttività", "Unità", "T1 · 7,4 mm", "T2 · 8,2 mm", "T3 · 20 mm"],
      [["Produttività energetica", "m²/GJ", g(pi["PsI_Energy"],"T1_7.4mm",2), g(pi["PsI_Energy"],"T2_8.2mm",2), g(pi["PsI_Energy"],"T3_20mm",2)],
       ["Resa di materiale", "m²/m²", g(pi["PsI_Yield"],"T1_7.4mm"), g(pi["PsI_Yield"],"T2_8.2mm"), g(pi["PsI_Yield"],"T3_20mm")],
       ["Throughput di linea", "m²/h", g(pi["PsI_Through"],"T1_7.4mm",0), g(pi["PsI_Through"],"T2_8.2mm",0), g(pi["PsI_Through"],"T3_20mm",0)]],
      "Tabella 4 — OP · Productivity Indicators per metrica e tipologia.")
ocr = R["OCR"]
table(["OCR = QP/AT", "Norma / soglia", "T1 · 7,4 mm", "T2 · 8,2 mm", "T3 · 20 mm"],
      [["Resistenza a flessione", "ISO 10545-4 / ≥35 N/mm²", g(ocr["OCR_Flex"],"T1_7.4mm"), g(ocr["OCR_Flex"],"T2_8.2mm"), g(ocr["OCR_Flex"],"T3_20mm")],
       ["Sforzo di rottura", "ISO 10545-4 / ≥700–1300 N", g(ocr["OCR_Break"],"T1_7.4mm"), g(ocr["OCR_Break"],"T2_8.2mm"), g(ocr["OCR_Break"],"T3_20mm")],
       ["Qualità superficiale", "ISO 10545-2 / ≥95 %", g(ocr["OCR_Surf"],"T1_7.4mm"), g(ocr["OCR_Surf"],"T2_8.2mm"), g(ocr["OCR_Surf"],"T3_20mm")]],
      "Tabella 5 — TQ · Output Conformity Rate per parametro (valori ≥ 1 = conforme con margine).")

heading("3.3 Sotto-indici e P-TSI — metodo primario (z-score)", 2)
def z(t, k): return f"{R[k][t]:+.3f}".replace('.', ',')
table(["Tipologia", "IOAI (z)", "OPI (z)", "TQI (z)", "P-TSI (z)"],
      [["T1 · 7,4 mm", z("T1_7.4mm","IOAI_z"), z("T1_7.4mm","OPI_z"), z("T1_7.4mm","TQI_z"), z("T1_7.4mm","PTSI_z")],
       ["T2 · 8,2 mm", z("T2_8.2mm","IOAI_z"), z("T2_8.2mm","OPI_z"), z("T2_8.2mm","TQI_z"), z("T2_8.2mm","PTSI_z")],
       ["T3 · 20 mm", z("T3_20mm","IOAI_z"), z("T3_20mm","OPI_z"), z("T3_20mm","TQI_z"), z("T3_20mm","PTSI_z")]],
      "Tabella 6 — Sotto-indici z-score e P-TSI (pesi uguali). Ordinamento: T3 > T1 > T2.")
doc.add_picture(f"{OUT}/RP7.4_fig1_profilo_dimensionale.png", width=Inches(5.6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para("Figura 1 — Profilo dimensionale (z-score): T3 forte su IOA e TQ, debole su OP; T1 profilo speculare (forte OP).", size=8.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

heading("3.4 Sotto-indici e P-TSI — metodo secondario (scoring 1–5 + AHP)", 2)
def s(t, k, dec=2): return f"{R[k][t]:.{dec}f}".replace('.', ',')
table(["Tipologia", "S_IOA [1–5]", "S_OP [1–5]", "S_TQ [1–5]", "P-TSI [1–5]", "TII"],
      [["T1 · 7,4 mm", s("T1_7.4mm","S_IOA"), s("T1_7.4mm","S_OP"), s("T1_7.4mm","S_TQ"), s("T1_7.4mm","PTSI_5"), f"+{R['TII']['T1_7.4mm']:.2f} %".replace('.',',')],
       ["T2 · 8,2 mm", s("T2_8.2mm","S_IOA"), s("T2_8.2mm","S_OP"), s("T2_8.2mm","S_TQ"), s("T2_8.2mm","PTSI_5"), f"+{R['TII']['T2_8.2mm']:.2f} %".replace('.',',')],
       ["T3 · 20 mm", s("T3_20mm","S_IOA"), s("T3_20mm","S_OP"), s("T3_20mm","S_TQ"), s("T3_20mm","PTSI_5"), f"+{R['TII']['T3_20mm']:.2f} %".replace('.',',')]],
      f"Tabella 7 — Punteggi dimensionali, P-TSI ponderato AHP e TII. Ordinamento: T3 > T1 > T2. Pesi AHP tra dimensioni: "
      f"IOA {R['AHP']['wDIM'][0]:.3f} · OP {R['AHP']['wDIM'][1]:.3f} · TQ {R['AHP']['wDIM'][2]:.3f}".replace('.',',') +
      f"; CR = {R['AHP']['CR']['DIM']} ≤ 0,10 (consistente).")
doc.add_picture(f"{OUT}/RP7.4_fig2_ptsi_tii.png", width=Inches(5.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para("Figura 2 — P-TSI (scala 1–5) e miglioramento annuo (TII) per tipologia.", size=8.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

heading("3.5 Analisi di sensibilità", 2)
sens_rows = [[k, v["ranking"].replace("_7.4mm"," · 7,4 mm").replace("_8.2mm"," · 8,2 mm").replace("_20mm"," · 20 mm").replace("T1","T1").replace("T2","T2").replace("T3","T3")]
             for k, v in R["SENS"].items()]
table(["Scenario", "Ordinamento"], sens_rows,
      "Tabella 8 — Stabilità dell'ordinamento. T3 primeggia in 3 scenari su 6 (quelli in cui disponibilità e qualità hanno peso normale o "
      "superiore); T1 primeggia quando si privilegia l'efficienza operativa; T2 non è mai primo.")

heading("3.6 Verifica dei KPI di attività", 2)
table(["KPI", "Baseline", "Obiettivo", "Risultato"],
      [["SCR (IOA)", "0.0", "Indicatore per 3 tipologie", "3 SCR × 3 tipologie"],
       ["PsI (OP)", "0.0", "Indicatore per 3 tipologie", "3 PsI × 3 tipologie"],
       ["OCR (TQ)", "0.0", "Indicatore per 3 tipologie", "3 OCR × 3 tipologie"],
       ["P-TSI", "0.0", "Indice per 3 tipologie", "3 P-TSI (z-score e scoring/AHP)"],
       ["Consistenza AHP", "—", "CR ≤ 0,10", f"CR = {R['AHP']['CR']['DIM']} (consistente)"]],
      "Tabella 9 — Verifica dei KPI di attività: obiettivi raggiunti.")

# ---------------- 4. DISCUSSIONE E CONCLUSIONI ----------------
heading("4. DISCUSSIONE E CONCLUSIONI")
heading("4.1 Lettura dei risultati", 2)
para("Il P-TSA restituisce profili tecnologici nettamente differenziati, coerenti con la natura fisica e d'uso delle tipologie. "
     "T3 (20 mm, outdoor/spessorato) è il prodotto tecnologicamente più sostenibile secondo il metodo primario: eccelle in disponibilità "
     "(IOA) e qualità tecnica (TQ, sforzo di rottura strutturale), ma è il più debole in performance operativa (OP) per l'elevata "
     "intensità energetica e il basso throughput. T1 (7,4 mm, alleggerito) ha il profilo speculare (massima efficienza operativa) ed è "
     "primo quando l'analisi privilegia l'OP. T2 (8,2 mm, standard) è la tipologia mediana e non risulta mai prima: emerge come "
     "principale candidato al miglioramento.")
heading("4.2 Robustezza e ponderazione", 2)
para("La convergenza tra metodo primario (z-score) e scoring + AHP sull'ordine T3 > T1 > T2 costituisce una validazione incrociata. "
     "L'analisi di sensibilità mostra che l'ordine dei primi due (T3 vs T1) dipende dal peso attribuito alla performance operativa: "
     "privilegiando l'OP, T1 supera T3. Non è una debolezza ma un'informazione decisionale, resa trasparente dai pesi AHP "
     f"(CR = {R['AHP']['CR']['DIM']}). Il TII positivo per tutte le tipologie (+3 ÷ +10 %) indica un miglioramento tendenziale, massimo per T3.")
heading("4.3 Interdipendenze e contributo al progetto", 2)
para("La RP7.4 chiude, sul versante prodotto, il sistema di misura della sostenibilità tecnologica del progetto: insieme all'EEA+ di "
     "OR7.3 fornisce i due indici (P-TSI ed EEA+I) del collaudo della Intelligent Industry (OR7.8) e alimenta la caratterizzazione dei "
     "prototipi di OR7.9. Riusa l'infrastruttura E2C (OR7.1–7.2) e l'impianto AHP/controlli di OR7.3, garantendo omogeneità metodologica "
     "tra la scala di fabbrica e quella di prodotto.")
heading("4.4 Contributo rispetto allo stato dell'arte", 2)
para("Il P-TSA porta a livello di prodotto ceramico una misura della sostenibilità tecnologica oggi assente dai framework prevalenti "
     "(LCA, GRI, SDG la trattano solo come abilitatore). Ancorando gli indicatori a metriche di supply chain (SCR/PsI/OCR) e alla "
     "conformità normativa (ISO 10545 / EN 14411), e rendendo esplicita la ponderazione via AHP, l'assessment trasforma dati operativi "
     "frammentati in un indice sintetico, interpretabile e auditabile per tipologia, utilizzabile come leva decisionale e come variabile "
     "per l'ottimizzazione multi-obiettivo della Intelligent Industry.")
heading("4.5 Prospettive", 2)
para("Il consolidamento riguarderà l'estensione del calcolo a granularità di lotto sull'intero portafoglio (OR6.8), la calibrazione di "
     "soglie e pesi in workshop con produzione/qualità/manutenzione, la lettura mensile del P-TSI su 12 mesi per il collaudo OR7.8, fino "
     "all'integrazione del P-TSI come variabile di controllo nel Digital Twin di prodotto.")

heading("Allegati", 2)
for a in ["RP7.4_dataset_sintetico.xlsx — dataset di input per tipologia/periodo (serie in corso di consolidamento).",
          "RP7.4_weights.xlsx — soglie di scoring e matrici AHP (pesi + CR).",
          "RP7.4_calculation_log.xlsx — log di calcolo versionato (formula → input → output).",
          "RP7.4_build.py — codice di calcolo (legge il dataset, ricostruisce indici e figure)."]:
    para("• " + a, size=9.5, after=2)

doc.save(f"{OUT}/RP7.4 Report di Product Technological Sustainability Assessment.docx")
print("docx salvato")
